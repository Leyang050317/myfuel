from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt


SOURCE = Path(r"C:\Users\Kimwen\Downloads\Appendix D - Task Description.docx")
OUTPUT = Path(r"C:\Users\Kimwen\Desktop\myfuel\Appendix D - Task Description - Er Kim Wen Completed.docx")


ANSWERS = {
    "Please briefly describe": [
        "I was mainly responsible for two areas of the MyFuel mobile application: (1) Fuel Price and Fuel Price Trend, and (2) Login and Logout.",
        "Fuel Price and Fuel Price Trend: I implemented the model-service-controller flow for retrieving Malaysia's current and historical RON95, RON97 and diesel prices. The service sends HTTP GET requests to the Malaysia Open Data API, decodes JSON into FuelPriceModel objects and filters historical level records. I also implemented the recent-history controller and the interactive multi-line trend chart with chronological dates, a legend and touch tooltips. The current prices and weekly changes are presented on the home page.",
        "Login and Logout: I implemented the login page with required-field validation, password visibility control, loading feedback and understandable authentication errors. Login uses Supabase email-and-password authentication and then retrieves the matching profile from the users table using the authenticated user's ID. The splash page checks for an existing session. For logout, the driver shell signs out the active Supabase session and replaces the current route with the login page.",
        "APIs used: Malaysia Open Data API (data.gov.my Data Catalogue, fuelprice dataset) for current and historical fuel-price records; Supabase Auth API for sign-in, sign-out and session management; and Supabase Database API for retrieving the user's profile.",
        "External libraries used: http ^1.2.1 for REST requests, fl_chart ^1.1.1 for the trend line chart, and supabase_flutter ^2.10.3 for authentication, session access and database queries. Flutter Material widgets and Navigator were used for the interface and page routing.",
    ],
    "What are the strengths": [
        "The modules use a clear separation of responsibilities. Models parse API data, services handle remote requests, controllers prepare information for the interface, and widgets focus on presentation. This makes the code easier to understand, maintain and extend.",
        "Fuel prices come from an official Malaysian government data source rather than being hard-coded. The trend chart compares three fuel types clearly, displays recent records from older to newer dates, and reveals exact values through touch tooltips.",
        "The authentication flow validates required fields, hides the password by default, disables repeated submissions while loading, and translates common authentication failures into clearer messages. Replacement navigation prevents users from simply returning to a protected screen after logout, while splash-screen session checking is convenient for returning users.",
    ],
    "What are the weaknesses": [
        "The fuel-price module depends on an internet connection and on the availability and format of the data.gov.my API. It currently has no retry mechanism, offline cache or explicit request timeout, and some screens do not provide dedicated loading, empty and error states. The current-price controller also assumes the API returns at least two records in the expected order.",
        "The trend view is limited to six recent level records and three fixed fuel types. Users cannot yet choose a date range, export data or compare additional categories.",
        "Although the repository parameter is named usernameOrEmail, the implemented Supabase login call accepts the value as an email address. Authentication and profile retrieval happen sequentially, so an authenticated account without a matching users-table row cannot complete login. Several repository methods are still unimplemented, and the fixed admin shortcut should eventually be replaced with unified role-based Supabase authentication.",
    ],
    "What have you learned": [
        "I learned how to connect a Flutter application to separate remote services while keeping networking, business logic and presentation responsibilities separated. I gained practical experience with Future, async/await, HTTP status checking, JSON decoding, typed Dart models and state updates after asynchronous loading.",
        "I also learned how Supabase manages authentication sessions, how to sign users in and out, and how to retrieve a profile using the authenticated user's ID. Using mounted checks taught me how to avoid updating the interface or navigating after a widget has been disposed.",
        "Building the trend chart improved my understanding of time-series transformation, chronological ordering, chart points, date labels, legends and tooltips. I also learned that validation, loading feedback, clear errors, responsive layouts and predictable navigation are essential parts of a reliable user experience.",
    ],
    "What are the challenges": [
        "One challenge was understanding the government fuel-price dataset. It contains different series types, so I had to identify level records for actual prices and arrange historical records in chronological order for the chart.",
        "Supabase integration was challenging because authentication data and application profile data are separate. After sign-in, I had to use the authenticated user's ID to query the users table and handle both authentication and database failures safely.",
        "Another challenge was coordinating asynchronous requests with Flutter's widget lifecycle. Loading flags and mounted checks were needed to prevent setState, messages or navigation after disposal. Presenting three fuel-price series on a small screen also required a compact legend, distinct colours, simplified date labels and touch tooltips.",
    ],
}


def insert_after(paragraph: Paragraph, text: str, answer_label: bool = False) -> Paragraph:
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    new_paragraph = Paragraph(node, paragraph._parent)
    new_paragraph.style = "Normal"
    new_paragraph.paragraph_format.space_before = Pt(0)
    new_paragraph.paragraph_format.space_after = Pt(6)
    new_paragraph.paragraph_format.line_spacing = 1.15

    if answer_label:
        label = new_paragraph.add_run("Answer: ")
        label.bold = True
        run = new_paragraph.add_run(text)
    else:
        run = new_paragraph.add_run(text)

    for current in new_paragraph.runs:
        current.font.name = "Arial"
        current.font.size = Pt(10)
        current._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
        current._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    return new_paragraph


def main() -> None:
    document = Document(SOURCE)
    questions = list(document.paragraphs)

    for question in questions:
        matching_key = next((key for key in ANSWERS if question.text.startswith(key)), None)
        if matching_key is None:
            continue

        cursor = question
        for index, answer in enumerate(ANSWERS[matching_key]):
            cursor = insert_after(cursor, answer, answer_label=index == 0)

    document.core_properties.title = "Appendix D - Task Description"
    document.core_properties.subject = "Individual contribution: Fuel Price, Fuel Price Trend, Login and Logout"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
